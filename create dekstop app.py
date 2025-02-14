import tkinter as tk
from tkinter import messagebox, filedialog, simpledialog, Toplevel, ttk
import ttkbootstrap as tb
from ttkbootstrap.constants import *
from ttkbootstrap.style import Style
from ttkbootstrap.tableview import Tableview  # ✅ Include Tableview
from tkcalendar import DateEntry
import sqlite3
import datetime
from datetime import datetime, timedelta
import os
import csv
import hashlib
import random
import subprocess
import platform
from PIL import Image, ImageTk
import pandas as pd
import base64  # ✅ Include base64
from Crypto.Cipher import AES, PKCS1_OAEP
from Crypto.PublicKey import RSA
from Crypto.Random import get_random_bytes
from cryptography.hazmat.primitives import serialization, asymmetric, hashes
from cryptography.hazmat.primitives.asymmetric import padding
from cryptography.hazmat.primitives.hashes import SHA256
from cryptography.hazmat.primitives.serialization import load_pem_private_key
from io import BytesIO
from Crypto.Util.Padding import pad, unpad
# Initialize the app with ttkbootstrap style
from key_request_function import request_access_window




# Initialize the app with a white theme
style = Style(theme="flatly")  # Start with a white-based theme
root = style.master
root.title("secureMSISDN")
root.iconbitmap('ooredooicon.ico')
root.geometry("500x500")

# Force pure white background
root.configure(background="white")





# Define the custom styles for buttons with specific size and no border for the red button
style.configure(
    "SmallRedButton.TButton",
    background="#d32f2f",          # Red background
    foreground="white",             # White text
    font=("Arial", 10, ),     # Smaller font size
    borderwidth=0,                  # No border
    focusthickness=0,               # No focus outline
)
style.configure(
    "WhiteButton.TButton",
    background="white", 
    foreground="#d32f2f", 
    font=("Arial", 10, ), 
    bordercolor="#d32f2f", 
    borderwidth=2
)



# Global variables
file_path = ""
files_data = []
current_user_department = ""
MASTER_KEY = os.getenv("MASTER_KEY") or b"16ByteMasterKey!"  # Example, must be securely stored

# Updated pool of security questions
SECURITY_QUESTIONS = [
    "What was the name of your first pet?",
    "What is your mother's maiden name?",
    "What is your favorite movie?",
    "What city were you born in?",
    "What was the make of your first car?"
]




    
# Password hashing
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

# Password reset with security question
def forgot_password():
    email = simpledialog.askstring("Forgot Password", "Enter your email:")
    
    if email:
        conn = sqlite3.connect('users.db')
        cursor = conn.cursor()
        cursor.execute('SELECT security_question, security_answer FROM users WHERE email = ?', (email,))
        user_data = cursor.fetchone()

        if user_data:
            stored_question, stored_answer = user_data

            # Ask the user their security question
            answer = simpledialog.askstring("Security Question", f"{stored_question}:")
            
            if answer and answer.lower() == stored_answer.lower():
                new_password = simpledialog.askstring("New Password", "Enter a new password:", show='*')
                hashed_password = hash_password(new_password)
                
                cursor.execute('UPDATE users SET password = ? WHERE email = ?', (hashed_password, email))
                conn.commit()
                messagebox.showinfo("Success", "Password reset successfully!")
            else:
                messagebox.showerror("Error", "Incorrect security answer.")
        else:
            messagebox.showerror("Error", "Email not found.")
        conn.close()

# Load RSA keys
def load_rsa_keys():
    with open("public.pem", "rb") as pub_file:
        public_key = RSA.import_key(pub_file.read())
    with open("private.pem", "rb") as priv_file:
        private_key = RSA.import_key(priv_file.read())
    return public_key, private_key







AES_KEY_SIZE = 32  # 256-bit AES keys

# Generate new RSA key pair
def generate_new_key():
    try:
        new_key = RSA.generate(2048)
        with open("private.pem", "wb") as priv_file:
            priv_file.write(new_key.export_key())
        with open("public.pem", "wb") as pub_file:
            pub_file.write(new_key.public_key().export_key())
        return "New key pair generated successfully."
    except Exception as e:
        return f"Error: {e}"

# Import an existing private key
def import_key(file_path):
    try:
        with open(file_path, "rb") as key_file:
            key = key_file.read()
        with open("private.pem", "wb") as priv_file:
            priv_file.write(key)
        return "Key imported successfully."
    except Exception as e:
        return f"Error: {e}"

# Export a specified RSA key
def export_key(key_type="private"):
    try:
        file_name = "private.pem" if key_type == "private" else "public.pem"
        with open(file_name, "rb") as key_file:
            return key_file.read()
    except Exception as e:
        return f"Error: {e}"

# View metadata of the private key
def view_key_metadata():
    try:
        file_stats = os.stat("private.pem")
        metadata = {
            "Key Type": "Private Key",
            "Size": file_stats.st_size,
            "Created On": datetime.fromtimestamp(file_stats.st_ctime),
            "Last Modified": datetime.fromtimestamp(file_stats.st_mtime),
        }
        return metadata
    except FileNotFoundError:
        return "No key found. Please generate or import a key."

# Inspect the current encryption key
def inspect_encryption_key():
    try:
        if not os.path.exists('encrypted_aes_key.bin'):
            raise FileNotFoundError("The encrypted AES key file ('encrypted_aes_key.bin') was not found.")

        with open('encrypted_aes_key.bin', 'rb') as f:
            encrypted_aes_key = f.read()

        with open('private.pem', 'rb') as priv_file:
            private_key = RSA.import_key(priv_file.read())

        cipher_rsa = PKCS1_OAEP.new(private_key)
        aes_key = cipher_rsa.decrypt(encrypted_aes_key)

        return aes_key
    except Exception as e:
        messagebox.showerror("Error", str(e))

# Regenerate and save a new AES encryption key
def regenerate_encryption_key():
    try:
        new_aes_key = os.urandom(AES_KEY_SIZE)
        with open('public.pem', 'rb') as pub_file:
            public_key = RSA.import_key(pub_file.read())

        cipher_rsa = PKCS1_OAEP.new(public_key)
        encrypted_new_aes_key = cipher_rsa.encrypt(new_aes_key)

        with open('encrypted_aes_key.bin', 'wb') as f:
            f.write(encrypted_new_aes_key)

        messagebox.showinfo("Success", "New AES key generated and saved for future encryption operations.")
    except Exception as e:
        messagebox.showerror("Error", f"Key regeneration failed: {e}")

# Key management window
def open_key_management_window():
    key_window = Toplevel()
    key_window.title("Key Management")
    key_window.geometry("400x300")

    inspect_btn = tk.Button(
        key_window,
        text="Inspect Encryption Key",
        command=lambda: messagebox.showinfo(
            "Key Details",
            f"Current Key: {base64.b64encode(inspect_encryption_key()).decode('utf-8')}"
        ),
        bg="red", fg="white", font=("Arial", 12, "bold")
    )
    inspect_btn.pack(pady=10)

    regenerate_btn = tk.Button(
        key_window,
        text="Regenerate Encryption Key",
        command=regenerate_encryption_key,
        bg="black", fg="white", font=("Arial", 12, "bold")
    )
    regenerate_btn.pack(pady=10)

    close_btn = tk.Button(
        key_window,
        text="Close",
        command=key_window.destroy,
        bg="gray", fg="white", font=("Arial", 12, "bold")
    )
    close_btn.pack(pady=10)

# Encrypt a given key using AES
def encrypt_key(key: bytes, master_key: bytes) -> str:
    try:
        cipher = AES.new(master_key, AES.MODE_ECB)
        encrypted = cipher.encrypt(pad(key, AES.block_size))
        return base64.b64encode(encrypted).decode('utf-8')
    except Exception as e:
        raise ValueError(f"Error encrypting key: {e}")

# Decrypt the key using AES
def decrypt_key(encrypted_key: str, master_key: bytes) -> bytes:
    try:
        cipher = AES.new(master_key, AES.MODE_ECB)
        decrypted = unpad(cipher.decrypt(base64.b64decode(encrypted_key)), AES.block_size)
        return decrypted
    except Exception as e:
        raise ValueError(f"Error decrypting key: {e}")





def add_logo(frame):
    try:
        # Open the logo with a transparent background
        logo_img = Image.open(r"C:\Users\youss\ooredoo_logo.png")
        
        # Resize proportionally based on app dimensions
        logo_width, logo_height = logo_img.size
        new_width = 200  # Adjust based on your layout
        new_height = int((new_width / logo_width) * logo_height)
        logo_img = logo_img.resize((new_width, new_height), Image.LANCZOS)
        
        # Convert to a Tkinter-compatible image
        logo_photo = ImageTk.PhotoImage(logo_img)
        logo_label = tb.Label(frame, image=logo_photo, background='white')  # Set background to match the app
        logo_label.image = logo_photo  # Keep a reference to avoid garbage collection
        logo_label.pack(pady=10)
        
    except Exception as e:
        print(f"Error loading logo: {e}")


    


    
    
# Ensure the cursor is available globally
def init_db():
    global conn, cursor
    conn = sqlite3.connect('app_data.db')
    cursor = conn.cursor()



def open_file(file_path):
    print(f"Debug: file_path just before use: {file_path}")
    print(f"Type: {type(file_path)}")
    if isinstance(file_path, bytes):
        print(f"Bytes content: {file_path[:100]}")

    try:
        if platform.system() == 'Windows':
            os.startfile(file_path)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', file_path))
        else:  # Linux
            subprocess.call(('xdg-open', file_path))
    except Exception as e:
        messagebox.showerror("Error", f"Unable to open file: {e}")

        


def log_file_activity(user_id, user_email, department, filename, file_size, key, activity_type="Encryption"):
    """
    Logs a file activity (encryption or decryption) into the file_activities table.
    
    Table Schema (file_activities):
    -------------------------------------------------
    id            INTEGER PRIMARY KEY AUTOINCREMENT
    user_id       INTEGER NOT NULL
    user_email    TEXT NOT NULL
    department    TEXT NOT NULL
    filename      TEXT NOT NULL
    file_size     INTEGER NOT NULL
    timestamp     TEXT NOT NULL
    activity_type TEXT NOT NULL
    key           TEXT
    -------------------------------------------------

    Parameters:
        user_id (int): The auto-increment PK of the user performing the activity.
        user_email (str): The email of the user.
        department (str): The department of the user.
        filename (str): The name of the file being encrypted or decrypted.
        file_size (int): The size of the file in bytes.
        key (str): The key used for encryption or decryption (will be encrypted before saving).
        activity_type (str): "Encryption" or "Decryption" (default: "Encryption").
    """
    try:
        # Connect to the SQLite database
        conn = sqlite3.connect('app_data.db')
        cursor = conn.cursor()

        # Get the current timestamp
        current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Encrypt the key before saving, assuming encrypt_key() and MASTER_KEY exist
        encrypted_key = encrypt_key(key.encode('utf-8'), MASTER_KEY)

        # Insert the activity log into the database
        cursor.execute('''
            INSERT INTO file_activities (
                user_id, user_email, department, filename, file_size, timestamp, activity_type, key
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (user_id, user_email, department, filename, file_size, current_time, activity_type, encrypted_key))

        # Commit changes and close connection
        conn.commit()
        conn.close()

    except Exception as e:
        print(f"Error logging file activity: {e}")








def open_document_encryption_tool():
    """
    Opens a tool for processing Word or Excel documents for encryption.
    """
    tool_window = tk.Toplevel()
    tool_window.title("Document Encryption Tool")
    tool_window.geometry("800x600")

    tk.Label(tool_window, text="Upload a Word or Excel file:", font=("Arial", 12)).pack(pady=10)
    file_path_var = tk.StringVar()

    # Frame to display content
    content_frame = tk.Frame(tool_window)
    content_frame.pack(fill="both", expand=True, padx=10, pady=10)

    def process_document(file_path):
        """
        Processes the uploaded document and displays its content.
        """
        # Debug: Inspect the file path or bytes content
        print(f"File path received: {file_path}")
        print(f"Type of file_path: {type(file_path)}")
        if isinstance(file_path, bytes):
            print(f"Bytes content (first 100 bytes): {file_path[:100]}")
            if not file_path.startswith(b"PK"):  # Check if bytes start with "PK"
                raise ValueError("Invalid Excel file content. File is not a valid .xlsx file.")

        # Process the file based on its extension
        try:
            if file_path.endswith(".xlsx"):
                display_excel_content(file_path, content_frame)
            elif file_path.endswith(".docx"):
                display_word_content(file_path, content_frame)
            else:
                raise ValueError("Unsupported file type. Please upload a .docx or .xlsx file.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process the file: {e}")

    
    def upload_document():
        """
        Handles document upload and reads the file content.
        """
        

        file_path = filedialog.askopenfilename(
            title="Select a Word or Excel File",
            filetypes=[("Word and Excel Files", "*.docx *.xlsx"), ("Word Documents", "*.docx"), ("Excel Files", "*.xlsx")]
        )

        if file_path:
            # Debug: Inspect the file path
            print(f"Selected file path: {file_path}")

            # Open the file in binary mode and read its content
            try:
                with open(file_path, "rb") as f:
                    file_content = f.read()  # Read the file content as bytes

                # Debug: Inspect the first few bytes of the file content
                print(f"File content (first 100 bytes): {file_content[:100]}")
                if not file_content.startswith(b"PK"):  # Validate the file is an Excel ZIP archive
                    raise ValueError("Invalid Excel file content. File is not a valid .xlsx file.")

                # Pass the bytes content to process_document
                process_document(file_path)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to read the file: {e}")


    

    tk.Button(tool_window, text="Upload Document", command=upload_document, font=("Arial", 10)).pack(pady=10)
    tk.Entry(tool_window, textvariable=file_path_var, state="readonly", width=60).pack(pady=5)

    tk.Button(tool_window, text="Return", command=tool_window.destroy, font=("Arial", 10)).pack(pady=10)



def display_excel_content(file_path, parent_frame):
    import pandas as pd

    try:
        # Load Excel data safely
        df = read_excel_safely(file_path)

        # Treeview for data preview (enable cell selection)
        tree = ttk.Treeview(parent_frame, selectmode="extended")  # Allow multiple selection
        tree.pack(side="left", fill="both", expand=True)

        # Scrollbar for Treeview
        tree_scroll = ttk.Scrollbar(parent_frame, orient="vertical", command=tree.yview)
        tree_scroll.pack(side="left", fill="y")
        tree.configure(yscrollcommand=tree_scroll.set)

        # Configure columns
        tree["columns"] = list(df.columns)
        tree["show"] = "headings"

        for col in df.columns:
            tree.heading(col, text=col)
            tree.column(col, width=100)

        # Insert rows
        for idx, row in df.iterrows():
            tree.insert("", "end", values=list(row))

        # Listbox to show selected cells
        listbox_frame = tk.Frame(parent_frame)
        listbox_frame.pack(side="right", fill="both", expand=True, padx=10)
        tk.Label(listbox_frame, text="Selected Cells:").pack(pady=5)

        selected_listbox = tk.Listbox(listbox_frame, height=15, width=40)
        selected_listbox.pack(fill="both", expand=True)

        # Store selected cells
        selected_cells = []

        # Function to handle cell selection
        def on_cell_click(event):
            region = tree.identify("region", event.x, event.y)
            if region == "cell":
                row_id = tree.identify_row(event.y)
                col_id = tree.identify_column(event.x)

                col_index = int(col_id.strip("#")) - 1
                col_name = df.columns[col_index]
                row_index = tree.index(row_id)

                cell = (row_index, col_name)
                if cell in selected_cells:
                    selected_cells.remove(cell)
                    for i, item in enumerate(selected_listbox.get(0, tk.END)):
                        if item == f"Row {row_index + 1}, Column '{col_name}'":
                            selected_listbox.delete(i)
                            break
                else:
                    selected_cells.append(cell)
                    selected_listbox.insert(tk.END, f"Row {row_index + 1}, Column '{col_name}'")

        # Bind click event
        tree.bind("<Button-1>", on_cell_click)

         # Button to remove selected cells manually
        def remove_selected_cells():
            selected_items = selected_listbox.curselection()
            if not selected_items:
                messagebox.showinfo("Info", "No items selected to remove.")
                return

            for index in reversed(selected_items):  # Iterate in reverse to avoid index shifting
                item_text = selected_listbox.get(index)
                # Parse row and column from the string (e.g., "Row 1, Column 'MSISDN'")
                row_col = item_text.split(", ")
                row_index = int(row_col[0].split(" ")[1]) - 1  # Extract row index
                col_name = row_col[1].split("'")[1]  # Extract column name
                selected_cells.remove((row_index, col_name))  # Remove from selected cells
                selected_listbox.delete(index)  # Remove from Listbox

        tk.Button(listbox_frame, text="Remove Selected Cells", command=remove_selected_cells).pack(pady=10)
        
        def encrypt_selected_cells():
            try:
                if not selected_cells:
                    messagebox.showerror("Error", "No cells selected for encryption.")
                    print("Debug: No cells selected for encryption.")
                    return

                # Create a copy of the DataFrame to avoid modifying the original file
                df_copy = df.copy()
                print("Debug: DataFrame copy created successfully.")

                # Generate AES key
                aes_key = get_random_bytes(16)
                print(f"Debug: Generated AES key (base64): {base64.b64encode(aes_key).decode()}")

                # Encrypt the selected cells
                for row_index, col_name in selected_cells:
                    if col_name not in df_copy.columns:
                        raise ValueError(f"Column '{col_name}' not found in the DataFrame.")

                    value = df_copy.at[row_index, col_name]  # Get the cell value

                    # Ensure column is cast to object to avoid type issues during encryption
                    if df_copy[col_name].dtype != 'object':
                        df_copy[col_name] = df_copy[col_name].astype('object')

                    if pd.notna(value):  # Only encrypt non-NaN values
                        encrypted_value = encrypt_data(value, aes_key)
                        df_copy.at[row_index, col_name] = encrypted_value
                        print(f"Debug: Encrypted cell at row {row_index}, column '{col_name}'. Original: {value}, Encrypted: {encrypted_value}")

                # Prompt user to save the encrypted file
                encrypted_file_path = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[("Excel files", "*.xlsx")],
                    title="Save Encrypted Excel File"
                )
                if not encrypted_file_path:
                    messagebox.showerror("Error", "File save operation was canceled.")
                    print("Debug: File save operation was canceled.")
                    return

                # Save the encrypted DataFrame to the specified file path
                df_copy.to_excel(encrypted_file_path, index=False)
                print(f"Debug: Encrypted file saved to {encrypted_file_path}")

                # Save the AES key encrypted with RSA
                aes_key_file = os.path.splitext(encrypted_file_path)[0] + "_aes_key.bin"
                try:
                    # Load RSA public key to encrypt AES key
                    with open("public.pem", "rb") as key_file:
                        public_key = serialization.load_pem_public_key(key_file.read())
                    encrypted_aes_key = public_key.encrypt(
                        aes_key,
                        padding.OAEP(mgf=padding.MGF1(algorithm=SHA256()), algorithm=SHA256(), label=None)
                    )

                    with open(aes_key_file, "wb") as f:
                        f.write(encrypted_aes_key)
                    print(f"Debug: AES key saved to {aes_key_file}")
                except Exception as e:
                    print(f"Debug: Failed to save AES key: {e}")
                    raise ValueError(f"Failed to save AES key: {e}")

                # Log the encryption activity
                log_file_activity(
                    user_id=current_user_id,
                    user_email=current_user_email,
                    department=current_user_department,
                    filename=os.path.basename(encrypted_file_path),
                    file_size=os.path.getsize(encrypted_file_path),
                    activity_type="Cell Encryption",
                    key=base64.b64encode(aes_key).decode('utf-8')  # Save the AES key in base64 format
                )
                print("Debug: Encryption activity logged successfully.")

                # Provide success feedback
                messagebox.showinfo(
                    "Success",
                    f"Encrypted file saved to {encrypted_file_path}\nAES key saved to {aes_key_file}"
                )
                print(f"Debug: Total cells encrypted: {len(selected_cells)}")

            except ValueError as ve:
                messagebox.showerror("Error", f"Validation error: {ve}")
                print(f"Debug: Validation error: {ve}")
            except Exception as e:
                messagebox.showerror("Error", f"Encryption failed: {e}")
                print(f"Debug: Encryption failed: {e}")




        # Add encryption button
        tk.Button(
            listbox_frame,
            text="Encrypt Selected Cells",
            command=encrypt_selected_cells
        ).pack(pady=10)

    except Exception as e:
        messagebox.showerror("Error", f"Failed to display Excel content: {e}")



def display_word_content(file_path, parent_frame):
    from docx import Document
    import tkinter as tk
    from tkinter import messagebox

    try:
        # Load Word document
        doc = Document(file_path)
        text_widget = tk.Text(parent_frame, wrap="word")
        text_widget.pack(fill="both", expand=True)

        # Insert document text into the Text widget
        for para in doc.paragraphs:
            text_widget.insert("end", para.text + "\n")

        # Cart to store selected text
        cart = []

        # Function to add selected text to the cart
        def add_to_cart():
            try:
                selected_text = text_widget.selection_get()
                if selected_text.strip():
                    cart.append(selected_text)
                    print(f"Debug: Added to cart: {selected_text}")
                    messagebox.showinfo("Cart", f"'{selected_text}' added to the cart.")
                else:
                    messagebox.showerror("Error", "No text selected. Please highlight text to add to the cart.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to add to cart: {e}")
                print(f"Debug: {e}")

        # Function to encrypt all items in the cart
        def encrypt_cart_items():
            if not cart:
                messagebox.showerror("Error", "The cart is empty. Please add items to the cart before encrypting.")
                return

            try:
                encrypt_selected_text_with_cart(text_widget, file_path, cart)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to encrypt cart items: {e}")
                print(f"Debug: {e}")

        # Add "Add to Cart" button
        tk.Button(
            parent_frame,
            text="Add to Cart",
            command=add_to_cart
        ).pack(pady=5)

        # Add "Encrypt Cart Items" button
        tk.Button(
            parent_frame,
            text="Encrypt Cart Items",
            command=encrypt_cart_items
        ).pack(pady=5)

    except Exception as e:
        messagebox.showerror("Error", f"Failed to load Word file: {e}")
        print(f"Debug: {e}")



def encrypt_selected_text_with_cart(text_widget, file_path, cart):
    """
    Encrypts selected text from a Word file and replaces it with encrypted text. Adds selected text to the cart.
    """
    from docx import Document
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
    from cryptography.hazmat.primitives.asymmetric import padding
    from cryptography.hazmat.primitives.hashes import SHA256
    from cryptography.hazmat.primitives import serialization
    import base64
    import os

    try:
        # Validate file_path
        if not isinstance(file_path, str) or not file_path.endswith(".docx"):
            raise ValueError("Invalid file_path. Expected a valid .docx file.")

        print(f"Debug: file_path before processing: {file_path}, Type: {type(file_path)}")

        # Verify that text_widget is a proper Text widget
        if not hasattr(text_widget, "tag_ranges"):
            raise ValueError("Invalid text_widget. Expected a valid Text widget.")

        # Get selected text ranges
        selected_ranges = text_widget.tag_ranges("sel")
        if not selected_ranges:
            messagebox.showerror("Error", "Please highlight text to add to the cart.")
            return

        # Extract selected text
        selected_text = text_widget.get(selected_ranges[0], selected_ranges[1]).strip()
        if not selected_text:
            messagebox.showerror("Error", "No valid text selected.")
            return

        # Add selected text to the cart
        cart.append(selected_text)
        print(f"Debug: Added to cart: {selected_text}")

        # Encrypt text in the cart
        if messagebox.askyesno("Encrypt", "Do you want to encrypt all items in the cart?"):
            aes_key = os.urandom(32)  # Use a 256-bit AES key
            iv = os.urandom(12)  # Recommended size for GCM is 12 bytes
            encrypted_text_map = {}

            # Encrypt each item in the cart
            for item in cart:
                cipher = Cipher(algorithms.AES(aes_key), modes.GCM(iv))
                encryptor = cipher.encryptor()
                ciphertext = encryptor.update(item.encode()) + encryptor.finalize()
                encrypted_text_map[item] = base64.b64encode(iv + encryptor.tag + ciphertext).decode()

            # Modify the Word file
            doc = Document(file_path)
            for para in doc.paragraphs:
                for original_text, encrypted_text in encrypted_text_map.items():
                    if original_text in para.text:
                        para.text = para.text.replace(original_text, encrypted_text)

            # Save the modified Word file
            updated_file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                title="Save Encrypted Word File"
            )
            if not updated_file_path:
                messagebox.showerror("Error", "File save operation was canceled.")
                return
            doc.save(updated_file_path)

            # Load RSA public key to encrypt AES key
            with open("public.pem", "rb") as key_file:
                public_key = serialization.load_pem_public_key(key_file.read())
            encrypted_aes_key = public_key.encrypt(
                aes_key,
                padding.OAEP(mgf=padding.MGF1(algorithm=SHA256()), algorithm=SHA256(), label=None)
            )

            # Save the encrypted AES key to a separate file
            aes_key_file = os.path.splitext(updated_file_path)[0] + "_aes_key.bin"
            with open(aes_key_file, "wb") as f:
                f.write(encrypted_aes_key)

            # Log activity
            log_file_activity(
                user_id=current_user_id,
                user_email=current_user_email,
                department=current_user_department,
                filename=os.path.basename(updated_file_path),
                file_size=os.path.getsize(updated_file_path),
                activity_type="Word File Encryption",
                key=base64.b64encode(aes_key).decode("utf-8")
            )

            # Display success message
            messagebox.showinfo(
                "Success",
                f"Encrypted Word file saved to {updated_file_path}\nAES key saved to {aes_key_file}"
            )
            print(f"Encrypted file saved at: {updated_file_path}")
            print(f"Key saved at: {aes_key_file}")

            # Clear the cart after encryption
            cart.clear()

    except ValueError as ve:
        messagebox.showerror("Error", f"Validation error: {ve}")
        print(f"Debug: {ve}")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")
        print(f"Debug: {e}")






from io import BytesIO
import pandas as pd

def read_excel_safely(file_path):
    """
    Safely reads an Excel file, handling both string paths and bytes objects.

    Parameters:
        file_path (str or BytesIO): Path to the Excel file or a BytesIO object.

    Returns:
        DataFrame: The loaded DataFrame.
    """
    try:
        if isinstance(file_path, BytesIO):
            print("Debug: Processing BytesIO input.")
            return pd.read_excel(file_path, engine='openpyxl')
        elif isinstance(file_path, str):
            print("Debug: Processing file path input.")
            return pd.read_excel(file_path, engine='openpyxl')
        else:
            raise ValueError("Invalid file_path type. Expected str or BytesIO.")
    except Exception as e:
        print(f"Debug: Failed to read the Excel file: {e}")
        raise ValueError(f"Failed to read the Excel file: {e}")



def validate_file_path(file_path):
    if isinstance(file_path, bytes):
        if not file_path.startswith(b"PK"):  # Excel files start with 'PK'
            with open("invalid_bytes_output.xlsx", "wb") as f:
                f.write(file_path)
            raise ValueError("Invalid bytes input. File saved for inspection.")
        return BytesIO(file_path)
    elif isinstance(file_path, str):
        if not file_path.endswith(('.xlsx', '.xls')):
            raise ValueError("Invalid file format. Only .xlsx and .xls files are supported.")
        return file_path
    else:
        raise ValueError("Invalid file_path type. Expected str or bytes.")









def encrypt_data(value, aes_key):
    """
    Encrypts a single value using the provided AES key.

    Parameters:
        value (str): The value to encrypt.
        aes_key (bytes): The AES key used for encryption.

    Returns:
        str: The encrypted value encoded in base64.
    """
    try:
        if pd.isna(value):  # Skip NaN values
            return value

        # Encrypt the value using AES
        cipher_aes = AES.new(aes_key, AES.MODE_EAX)
        nonce = cipher_aes.nonce
        ciphertext, tag = cipher_aes.encrypt_and_digest(str(value).encode('utf-8'))

        # Combine nonce, tag, and ciphertext into a single base64-encoded string
        encrypted_value = base64.b64encode(nonce + tag + ciphertext).decode('utf-8')

        print(f"Debug: Successfully encrypted value: {value} -> {encrypted_value}")
        return encrypted_value

    except Exception as e:
        print(f"Debug: Encryption failed for value '{value}': {e}")
        raise ValueError(f"Encryption failed for value '{value}': {e}")






def upload_file():
    """
    Prompts the user to upload an Excel or Word file and sets the global file_path variable.
    """
    global file_path
    file_path = filedialog.askopenfilename(
        filetypes=[
            ("Supported Files", "*.xlsx *.xls *.docx"),
            ("Excel Files", "*.xlsx *.xls"),
            ("Word Files", "*.docx"),
        ],
        title="Select a file to decrypt",
    )

    if not file_path:
        messagebox.showerror("Error", "No file selected.")
        return

    # Check the file extension
    file_extension = os.path.splitext(file_path)[-1].lower()
    if file_extension not in [".xlsx", ".xls", ".docx"]:
        messagebox.showerror("Error", "Unsupported file type. Please upload an Excel or Word file.")
        file_path = None  # Reset the file_path
        return

    # Provide feedback to the user
    messagebox.showinfo("Success", f"File selected: {os.path.basename(file_path)}")




def decrypt_data():
    """
    Decrypts either an Excel or Word file based on the file type.
    """
    try:
        if not file_path:
            messagebox.showerror("Error", "Please upload an encrypted file first.")
            return

        print(f"Debug: Uploaded file_path: {file_path}")

        file_extension = os.path.splitext(file_path)[-1].lower()
        key_file_path = os.path.splitext(file_path)[0] + "_aes_key.bin"
        print(f"Debug: File extension: {file_extension}")
        print(f"Debug: Associated AES key file: {key_file_path}")

        if not os.path.exists(key_file_path):
            messagebox.showerror("Error", f"AES key file not found for {os.path.basename(file_path)}.")
            return

        try:
            with open("private.pem", "rb") as f:
                private_key = serialization.load_pem_private_key(f.read(), password=None)
            print(f"Debug: Private RSA key loaded successfully: {private_key.key_size} bits")

        except Exception as e:
            print(f"Debug: Failed to load private RSA key: {e}")
            raise

        try:
            with open(key_file_path, "rb") as f:
                encrypted_aes_key = f.read()
            print(f"Debug: Encrypted AES key loaded (length: {len(encrypted_aes_key)}).")
        except Exception as e:
            print(f"Debug: Failed to read encrypted AES key file: {e}")
            raise

        print(f"Debug: Encrypted AES key (base64): {base64.b64encode(encrypted_aes_key).decode()}")

        try:
            aes_key = private_key.decrypt(
                encrypted_aes_key,
                padding.OAEP(
                    mgf=padding.MGF1(algorithm=SHA256()),
                    algorithm=SHA256(),
                    label=None
                )
            )
            print(f"Debug: Decrypted AES key (base64): {base64.b64encode(aes_key).decode()}")
        except Exception as e:
            print(f"Debug: Failed to decrypt AES key: {e}")
            raise

        if file_extension in [".xlsx", ".xls"]:
            decrypt_excel(file_path, aes_key)
        elif file_extension == ".docx":
            decrypt_word(file_path, aes_key)
        else:
            messagebox.showerror("Error", "Unsupported file type. Please upload an Excel or Word file.")
            return

        log_file_activity(
            user_id=current_user_id,
            user_email=current_user_email,
            department=current_user_department,
            filename=os.path.basename(file_path),
            file_size=os.path.getsize(file_path),
            activity_type="Decryption",
            key=base64.b64encode(aes_key).decode("utf-8"),
        )
        print("Debug: Decryption process completed successfully.")

    except Exception as e:
        messagebox.showerror("Error", f"Decryption failed: {e}")
        print(f"Debug: Decryption failed: {e}")



def decrypt_excel(file_path, aes_key):
    """
    Decrypts an Excel file using the provided AES key.

    Parameters:
        file_path (str): Path to the encrypted Excel file.
        aes_key (bytes): The AES key to decrypt the file.

    Raises:
        ValueError: If decryption fails or input parameters are invalid.
    """
    try:
        # Validate input parameters
        if not isinstance(file_path, str) or not file_path.endswith((".xlsx", ".xls")):
            raise ValueError("Invalid file_path. Must be a valid .xlsx or .xls file.")
        if not isinstance(aes_key, bytes) or len(aes_key) not in [16, 24, 32]:
            raise ValueError("Invalid AES key. Ensure it's a 128-bit, 192-bit, or 256-bit key.")

        # Debug: Log the file path and key details
        print(f"Debug: Decrypting file: {file_path}")
        print(f"Debug: AES key (base64): {base64.b64encode(aes_key).decode()}")

        # Load the encrypted Excel file
        df = pd.read_excel(file_path)
        print("Debug: Encrypted Excel file loaded successfully.")

        # Function to decrypt a single cell
        def decrypt_cell(value):
            if pd.isna(value) or not is_base64_encoded(value):
                print(f"Debug: Skipping non-encrypted value: {value}")
                return value  # Skip non-encrypted values
            try:
                print(f"Debug: Decrypting value: {value}")
                data = base64.b64decode(value)
                nonce, tag, ciphertext = data[:16], data[16:32], data[32:]
                cipher_aes = AES.new(aes_key, AES.MODE_EAX, nonce=nonce)
                decrypted_value = cipher_aes.decrypt_and_verify(ciphertext, tag).decode('utf-8')
                print(f"Debug: Decrypted value: {decrypted_value}")
                return decrypted_value
            except Exception as decryption_error:
                print(f"Debug: Failed to decrypt value: {value}. Error: {decryption_error}")
                raise ValueError(f"Decryption failed for value: {value}. Error: {decryption_error}")

        # Decrypt all cells in the DataFrame
        print("Debug: Starting decryption of all cells.")
        for col in df.columns:
            print(f"Debug: Decrypting column: {col}")
            df[col] = df[col].apply(decrypt_cell)

        # Save the decrypted file
        decrypted_file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")],
            title="Save Decrypted Excel File"
        )
        if not decrypted_file_path:
            raise ValueError("Decryption canceled. No save path specified.")

        df.to_excel(decrypted_file_path, index=False)
        messagebox.showinfo("Success", f"Decrypted Excel file saved to {decrypted_file_path}")
        print(f"Debug: Decrypted file saved to: {decrypted_file_path}")

    except ValueError as ve:
        messagebox.showerror("Error", f"Validation error: {ve}")
        print(f"Debug: {ve}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to decrypt Excel file: {e}")
        print(f"Debug: Failed to decrypt Excel file: {e}")



import re

def is_base64_encoded(value):
    if not isinstance(value, str):
        return False
    try:
        decoded = base64.b64decode(value, validate=True)
        # Optionally, enforce a minimum length for encrypted data
        return len(decoded) > 32  # Minimum length for nonce + tag + ciphertext
    except Exception:
        return False


from docx import Document
def decrypt_word(file_path, aes_key):
    """
    Decrypts a Word file using the provided AES key, processing each word individually.
    """
    from docx import Document
    import base64
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
    from tkinter import filedialog, messagebox

    try:
        doc = Document(file_path)
        print(f"Debug: Loaded Word file: {file_path}")

        def decrypt_word_text(encrypted_word):
            """
            Decrypts a single encrypted word using AES-GCM.
            """
            if not is_base64_encoded(encrypted_word):
                print(f"Debug: Skipping non-encrypted word: {encrypted_word}")
                return encrypted_word  # Skip non-encrypted words

            try:
                data = base64.b64decode(encrypted_word)
                iv, tag, ciphertext = data[:12], data[12:28], data[28:]
                cipher = Cipher(algorithms.AES(aes_key), modes.GCM(iv))
                decryptor = cipher.decryptor()
                decrypted_word = decryptor.update(ciphertext) + decryptor.finalize_with_tag(tag)
                return decrypted_word.decode("utf-8")
            except Exception as e:
                print(f"Debug: Failed to decrypt word: {encrypted_word}. Error: {e}")
                return encrypted_word  # If decryption fails, return the original word

        # Process each paragraph word by word
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                print(f"Debug: Original paragraph text: {para.text}")
                words = para.text.split()  # Split paragraph into individual words
                decrypted_words = [decrypt_word_text(word) for word in words]
                para.text = " ".join(decrypted_words)  # Reconstruct the paragraph
                print(f"Debug: Decrypted paragraph text: {para.text}")

        # Save the decrypted Word file
        decrypted_file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            title="Save Decrypted Word File"
        )
        if not decrypted_file_path:
            messagebox.showerror("Error", "File save operation was canceled.")
            return

        doc.save(decrypted_file_path)
        messagebox.showinfo("Success", f"Decrypted Word file saved to {decrypted_file_path}")
        print(f"Decrypted Word file saved at: {decrypted_file_path}")

    except Exception as e:
        messagebox.showerror("Error", f"Failed to decrypt Word file: {e}")
        print(f"Debug: Failed to decrypt Word file: {e}")








# Function to simulate cloud upload
def simulate_upload():
    messagebox.showinfo("Cloud Upload", "Simulated upload to cloud server completed!")

def view_all_users():
    global root
    for widget in root.winfo_children():
        widget.destroy()

    user_frame = tb.Frame(root, padding=(20, 10), bootstyle="secondary")
    user_frame.pack(fill="both", expand=True)

    tb.Label(user_frame, text="All Registered Users", bootstyle="info").pack(pady=10)

    # Fetch all users, including user_id, first_name, last_name, and email
    users = get_all_users()  # This should return a list of dicts like [{"user_id": 1, "first_name": ..., ...}, ...]

    # Updated columns to include user_id
    columns = ('user_id', 'first_name', 'last_name', 'email')
    
    # Prepare row data for the Tableview
    row_data = []
    for user in users:
        row_data.append((
            user['user_id'],
            user['first_name'],
            user['last_name'],
            user['email']
        ))

    user_table = Tableview(
        master=user_frame,
        coldata=columns,
        rowdata=row_data,
        bootstyle="info"
    )
    user_table.pack(pady=10, fill='both', expand=True)

    # Button to remove users
    tb.Button(
        user_frame,
        text="Remove Selected User",
        command=lambda: remove_selected_user(user_table),
        bootstyle="danger-outline"
    ).pack(pady=10)

    # Return button to go back to the dashboard
    tb.Button(
        user_frame,
        text="Return",
        command=show_dashboard,
        bootstyle="secondary-outline"
    ).pack(pady=10)

def remove_user(user_id):
    """
    Deletes a user from the 'users' table using the user's numeric ID.
    
    :param user_id: The auto-incremented ID of the user to be removed.
    """
    try:
        conn = sqlite3.connect('users.db')
        cursor = conn.cursor()
        
        # Delete the user by user_id
        cursor.execute("DELETE FROM users WHERE user_id = ?", (user_id,))
        conn.commit()
        conn.close()

        # Optionally, refresh or return to the dashboard after removal
        show_dashboard()

    except Exception as e:
        messagebox.showerror("Error", f"Failed to remove user with ID {user_id}. Details: {e}")


def remove_selected_user(table):
    # Get the selected row; each row is a tuple of (user_id, first_name, last_name, email)
    selected = table.get_selected()
    if selected:
        # user_id is in the first column (index 0) after the update
        selected_user_id = selected[0][0]
        remove_user(selected_user_id)  # Updated remove_user to accept user_id
        messagebox.showinfo("Success", f"User with ID {selected_user_id} has been removed.")
    else:
        messagebox.showerror("Error", "No user selected.")


def manage_key_requests_window():
    req_win = tk.Toplevel()
    req_win.title("Key Requests")

    tree = ttk.Treeview(req_win, columns=("ReqID","UserID","File","Reason","Status","Time"), show="headings")
    tree.heading("ReqID", text="Request ID")
    tree.heading("UserID", text="User ID")
    tree.heading("File", text="File Name")
    tree.heading("Reason", text="Reason")
    tree.heading("Status", text="Status")
    tree.heading("Time", text="Request Time")
    tree.pack(fill="both", expand=True)

    def refresh_requests():
        tree.delete(*tree.get_children())
        conn = sqlite3.connect("ooredoo.db")
        cursor = conn.cursor()
        cursor.execute("""
            SELECT request_id, user_id, file_name, request_reason, status, request_timestamp
            FROM key_requests
            WHERE status='PENDING'
        """)
        rows = cursor.fetchall()
        conn.close()
        for row in rows:
            tree.insert("", tk.END, values=row)

    def approve_request():
        selected = tree.selection()
        if not selected:
            return
        values = tree.item(selected[0])["values"]
        request_id = values[0]

        conn = sqlite3.connect("ooredoo.db")
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE key_requests
            SET status='APPROVED', admin_id=?, approval_timestamp=datetime('now')
            WHERE request_id=?
        """, (current_admin_id, request_id))  # admin_id is the logged-in admin
        conn.commit()
        conn.close()

        messagebox.showinfo("Info", f"Request {request_id} approved.")
        refresh_requests()

    def reject_request():
        selected = tree.selection()
        if not selected:
            return
        values = tree.item(selected[0])["values"]
        request_id = values[0]

        conn = sqlite3.connect("ooredoo.db")
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE key_requests
            SET status='REJECTED', admin_id=?, approval_timestamp=datetime('now')
            WHERE request_id=?
        """, (current_admin_id, request_id))
        conn.commit()
        conn.close()

        messagebox.showinfo("Info", f"Request {request_id} rejected.")
        refresh_requests()

    btn_frame = tk.Frame(req_win)
    btn_frame.pack(pady=5)
    approve_btn = tk.Button(btn_frame, text="Approve", command=approve_request)
    reject_btn = tk.Button(btn_frame, text="Reject", command=reject_request)
    approve_btn.grid(row=0, column=0, padx=5)
    reject_btn.grid(row=0, column=1, padx=5)
     # Add a "Return" button
    def return_to_dashboard():
        req_win.destroy()  # Close the current window
        show_dashboard()   # Redirect to the admin dashboard

    return_btn = tk.Button(
        req_win,
        text="Return",
        command=return_to_dashboard
    )
    return_btn.pack(pady=10)
    
    refresh_requests()

def return_to_operation_choice(window):
    window.destroy()  # Close current window
    create_operation_choice_page()  # Reopen the main operation choice window

def show_dashboard():
    global current_user_department, last_page

    # Check if the logged-in user is admin
    if current_user_department != 'Admin':
        messagebox.showerror("Access Denied", "Only admins can access the dashboard.")
        return

    # Clear previous frame
    for widget in root.winfo_children():
        widget.destroy()

    dashboard_frame = tb.Frame(root, padding=(20, 10), bootstyle="secondary")
    dashboard_frame.pack(fill="both", expand=True)

    # Title for the dashboard
    tb.Label(
        dashboard_frame, 
        text="Admin Dashboard - File Activities", 
        bootstyle="info", 
        font=("Arial", 18, "bold")
    ).pack(pady=10)

    try:
        # Fetch file activities from the database
        conn = sqlite3.connect('app_data.db')
        cursor = conn.cursor()
        # Now also select user_id from file_activities if it exists
        cursor.execute("""
            SELECT 
                user_id, 
                user_email, 
                department, 
                filename, 
                file_size, 
                timestamp, 
                activity_type,
                key 
                
            FROM file_activities
        """)
        file_activities = cursor.fetchall()

        # Process keys to obfuscate them for security (show only last 4 chars if length >= 4)
        # We'll also keep user_id in the row data so we can display it
        processed_activities = []
        for uid, email, dept, fname, fsize, ts,  activity,key in file_activities:
            if key and len(key) >= 4:
                obfuscated_key = f"******{key[-4:]}"
            else:
                obfuscated_key = "N/A"
            processed_activities.append((
                uid,
                email,
                dept,
                fname,
                fsize,
                ts,
                activity,
                obfuscated_key,
            ))

    except Exception as e:
        messagebox.showerror("Database Error", f"Failed to fetch file activities: {e}")
        return
    finally:
        if conn:
            conn.close()

    # Table Columns (now includes 'User ID')
    columns = (
        'User ID', 
        'User Email', 
        'Department', 
        'File Name', 
        'File Size (Bytes)', 
        'Timestamp', 
        
        'Activity Type',
        'Key'
    )

    # Create TableView
    file_table = Tableview(
        master=dashboard_frame,
        coldata=columns,
        rowdata=processed_activities,
        paginated=True,
        searchable=True,
        bootstyle="success"
    )
    file_table.pack(pady=10, fill='both', expand=True)

    # View All Users button
    tb.Button(
        dashboard_frame, 
        text="View All Users", 
        command=view_all_users, 
        bootstyle="secondary-outline"
    ).pack(pady=10)

    # Key Management button
    tb.Button(
        dashboard_frame,
        text="Key Management",
        command=open_key_management_window,
        bootstyle="warning-outline"
    ).pack(pady=10)

    # Manage Key Requests button (admin feature)
    manage_requests_button = tk.Button(dashboard_frame, text="Manage Key Requests", command=manage_key_requests_window)
    manage_requests_button.pack(pady=10)

    # Dynamic Return button
    def navigate_back():
        if last_page == "encryption":
            create_main_page(is_encryption=True)
        elif last_page == "decryption":
            create_main_page(is_encryption=False)

    tb.Button(
        dashboard_frame,
        text="Return",
        command=lambda: return_to_operation_choice(dashboard_frame),
        bootstyle="secondary-outline"
    ).pack(pady=10)


# Fetch all users from SQLite (unchanged, except your final code might include user_id as well)
def get_all_users():
    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()
    cursor.execute("SELECT first_name, last_name, email FROM users")
    users = cursor.fetchall()


# Call init_db() when the app starts to ensure the database is ready
init_db()


# ------------------- Equipment Management Window -------------------




def manage_equipment():
    equipment_window = tb.Toplevel(root)
    equipment_window.title("Equipment Management")
    equipment_window.geometry("900x600")

    tb.Label(equipment_window, text="Equipment Management", font=("Arial", 14, "bold"), bootstyle="primary").pack(pady=10)

    # Frame for Table and Scrollbars
    table_frame = tb.Frame(equipment_window)
    table_frame.pack(fill="both", expand=True, padx=10, pady=10)

    # Define Columns
    columns = ("ID", "Name", "Category", "Manufacturer", "Model", "Serial Number",
               "Purchase Date", "Warranty End Date", "Status", "Location", "Assigned Project")

    # Treeview with Scrollbars
    equipment_table = tb.Treeview(table_frame, columns=columns, show="headings", bootstyle="info")
    for col in columns:
        equipment_table.heading(col, text=col)
        equipment_table.column(col, width=120, anchor="center")

    # Scrollbars
    y_scroll = tb.Scrollbar(table_frame, orient="vertical", command=equipment_table.yview)
    x_scroll = tb.Scrollbar(table_frame, orient="horizontal", command=equipment_table.xview)
    equipment_table.configure(yscrollcommand=y_scroll.set, xscrollcommand=x_scroll.set)
    y_scroll.pack(side="right", fill="y")
    x_scroll.pack(side="bottom", fill="x")
    equipment_table.pack(fill="both", expand=True)

    # ----------------- Search and Filter Section -----------------
    filter_frame = tb.Frame(equipment_window)
    filter_frame.pack(fill="x", padx=10, pady=5)

    tb.Label(filter_frame, text="Search:", bootstyle="info").pack(side="left", padx=5)
    search_entry = tb.Entry(filter_frame)
    search_entry.pack(side="left", padx=5)

    def fetch_categories():
        with sqlite3.connect("ooredoo_equipements.db") as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM Categories")
            categories = ["All"] + [row[0] for row in cursor.fetchall()]
        return categories

    # Dropdowns
    def update_dropdowns():
        category_menu["menu"].delete(0, "end")
        status_menu["menu"].delete(0, "end")

        for category in fetch_categories():
            category_menu["menu"].add_command(label=category, command=lambda value=category: category_var.set(value))

        for status in ["All", "Active", "In Maintenance", "Decommissioned"]:
            status_menu["menu"].add_command(label=status, command=lambda value=status: status_var.set(value))

    tb.Label(filter_frame, text="Category:", bootstyle="info").pack(side="left", padx=5)
    category_var = tb.StringVar(value="All")
    category_menu = tb.OptionMenu(filter_frame, category_var, *fetch_categories())
    category_menu.pack(side="left", padx=5)

    tb.Label(filter_frame, text="Status:", bootstyle="info").pack(side="left", padx=5)
    status_var = tb.StringVar(value="All")
    status_menu = tb.OptionMenu(filter_frame, status_var, "All", "Active", "In Maintenance", "Decommissioned")
    status_menu.pack(side="left", padx=5)

    tb.Button(filter_frame, text="Apply Filters",
              command=lambda: load_equipment(search_entry.get(), category_var.get(), status_var.get()),
              bootstyle="primary").pack(side="left", padx=5)

    # ----------------- Load Equipment Data -----------------
    def load_equipment(search_query="", category_filter="All", status_filter="All"):
        equipment_table.delete(*equipment_table.get_children())

        with sqlite3.connect("ooredoo_equipements.db") as conn:
            cursor = conn.cursor()
            query = """
                SELECT Equipments.id, Equipments.name, Categories.name, Equipments.manufacturer, Equipments.model, 
                    Equipments.serial_number, Equipments.purchase_date, Equipments.warranty_end_date, Equipments.status, 
                    Equipments.location, Equipments.assigned_project 
                FROM Equipments
                JOIN Categories ON Equipments.category_id = Categories.id
                WHERE 1=1
            """
            params = []
            if search_query.strip():
                query += " AND (Equipments.name LIKE ? OR Equipments.manufacturer LIKE ? OR Equipments.serial_number LIKE ?)"
                params.extend([f"%{search_query}%", f"%{search_query}%", f"%{search_query}%"])

            if category_filter != "All":
                query += " AND Categories.name = ?"
                params.append(category_filter)

            if status_filter != "All":
                query += " AND Equipments.status = ?"
                params.append(status_filter)

            cursor.execute(query, params)
            results = cursor.fetchall()
            for row in results:
                equipment_table.insert("", "end", values=row)

        update_dropdowns()

    # ----------------- Schedule Maintenance -----------------
    def schedule_maintenance():
        selected_item = equipment_table.selection()
        if not selected_item:
            messagebox.showerror("Error", "Please select an equipment to schedule maintenance.")
            return

        item_data = equipment_table.item(selected_item)["values"]
        equip_id = item_data[0]  # Get Equipment ID

        schedule_window = tb.Toplevel(equipment_window)
        schedule_window.title("Schedule Maintenance")
        schedule_window.geometry("400x300")

        # Frame for layout
        frame = tb.Frame(schedule_window)
        frame.pack(padx=10, pady=10, fill="both", expand=True)

        # Corrected Date Picker (Fix conflicts)
        tb.Label(frame, text="Select Maintenance Date:", bootstyle="info").pack(pady=10)
        maintenance_date = DateEntry(frame, date_pattern="yyyy-mm-dd", width=12, background="white", foreground="black", borderwidth=2)
        maintenance_date.pack(pady=5, fill="x")

        # Notes Entry (Optional)
        tb.Label(frame, text="Notes (Optional):", bootstyle="info").pack(pady=5)
        notes_entry = tb.Entry(frame)
        notes_entry.pack(pady=5, fill="x")

        def save_maintenance():
            selected_date = maintenance_date.get_date().strftime("%Y-%m-%d")

            with sqlite3.connect("ooredoo_equipements.db") as conn:
                cursor = conn.cursor()

                # Check if maintenance is already scheduled
                cursor.execute("""
                    SELECT COUNT(*) FROM MaintenanceSchedule 
                    WHERE equipment_id = ? AND maintenance_date = ? AND status = 'Scheduled'
                """, (equip_id, selected_date))
                existing_maintenance = cursor.fetchone()[0]

                if existing_maintenance > 0:
                    messagebox.showerror("Error", "Maintenance is already scheduled for this date.")
                    return

                # Insert maintenance record
                cursor.execute("""
                    INSERT INTO MaintenanceSchedule (equipment_id, maintenance_date, status, notes) 
                    VALUES (?, ?, 'Scheduled', ?)""",
                    (equip_id, selected_date, notes_entry.get()))
                
                # Update status to "In Maintenance"
                cursor.execute("UPDATE Equipments SET status='In Maintenance' WHERE id=?", (equip_id,))
                conn.commit()

            schedule_window.destroy()
            load_equipment()
            messagebox.showinfo("Success", "Maintenance Scheduled Successfully!")

        tb.Button(frame, text="Schedule Maintenance", command=save_maintenance, bootstyle="success").pack(pady=10)

        # ----------------- Buttons -----------------
    
    tb.Button(equipment_window, text="Edit Equipment", bootstyle="warning-outline").pack(pady=5)
    tb.Button(equipment_window, text="Delete Equipment", bootstyle="danger-outline").pack(pady=5)
    tb.Button(equipment_window, text="Refresh List", command=load_equipment, bootstyle="info-outline").pack(pady=5)
    tb.Button(equipment_window, text="Return", bootstyle="danger-outline").pack(pady=10)
    tb.Button(equipment_window, text="Schedule Maintenance", command=schedule_maintenance, bootstyle="warning-outline").pack(pady=5)

    load_equipment()















# ------------------- Project Management Window -------------------
def manage_projects():
    project_window = tb.Toplevel(root)
    project_window.title("Project Management")
    project_window.geometry("700x500")

    tb.Label(project_window, text="Project Management", font=("Arial", 14, "bold"), bootstyle="primary").pack(pady=20)

    # Table to display projects
    columns = ("ID", "Project Name", "Status", "Manager", "Start Date", "End Date")
    project_table = tb.Treeview(project_window, columns=columns, show="headings", bootstyle="info")

    for col in columns:
        project_table.heading(col, text=col)
        project_table.column(col, width=120)

    project_table.pack(fill="both", expand=True, padx=20, pady=10)

    # Function to load project data
    def load_projects():
        project_table.delete(*project_table.get_children())  # Clear table
        conn = sqlite3.connect("ooredoo.db")
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM projects")
        for row in cursor.fetchall():
            project_table.insert("", "end", values=row)
        conn.close()

    # Function to add a new project
    def add_project():
        add_window = tb.Toplevel(project_window)
        add_window.title("Add Project")
        add_window.geometry("400x350")

        tb.Label(add_window, text="Project Name:").pack(pady=5)
        entry_name = tb.Entry(add_window)
        entry_name.pack(pady=5)

        tb.Label(add_window, text="Status:").pack(pady=5)
        entry_status = tb.Entry(add_window)
        entry_status.pack(pady=5)

        tb.Label(add_window, text="Manager:").pack(pady=5)
        entry_manager = tb.Entry(add_window)
        entry_manager.pack(pady=5)

        tb.Label(add_window, text="Start Date (YYYY-MM-DD):").pack(pady=5)
        entry_start_date = tb.Entry(add_window)
        entry_start_date.pack(pady=5)

        tb.Label(add_window, text="End Date (YYYY-MM-DD):").pack(pady=5)
        entry_end_date = tb.Entry(add_window)
        entry_end_date.pack(pady=5)

        def save_project():
            conn = sqlite3.connect("ooredoo.db")
            cursor = conn.cursor()
            cursor.execute("INSERT INTO projects (name, status, manager, start_date, end_date) VALUES (?, ?, ?, ?, ?)",
                           (entry_name.get(), entry_status.get(), entry_manager.get(), entry_start_date.get(), entry_end_date.get()))
            conn.commit()
            conn.close()
            add_window.destroy()
            load_projects()

        tb.Button(add_window, text="Add Project", command=save_project, bootstyle="success").pack(pady=10)

    tb.Button(project_window, text="Add Project", command=add_project, bootstyle="success-outline").pack(pady=10)
    tb.Button(project_window, text="Refresh List", command=load_projects, bootstyle="info-outline").pack(pady=5)

     # 🔙 Return Button
    tb.Button(project_window, text="Return", command=lambda: return_to_operation_choice(project_window), bootstyle="danger-outline").pack(pady=10)
    load_projects()  # Load projects initially



def create_main_page(is_encryption=True):
    global current_user_department, last_page

    # Set last_page based on the current mode
    last_page = "encryption" if is_encryption else "decryption"

    # Clear previous widgets
    for widget in root.winfo_children():
        widget.destroy()

    # Create the main frame
    main_frame = tb.Frame(root, padding=(20, 10), bootstyle="white")
    main_frame.pack(fill="both", expand=True)

    # Add the app logo
    add_logo(main_frame)

    # Buttons for encryption mode
    if is_encryption:
        
        tb.Button(main_frame,text="Encrypt Sensitive Data in Documents",command=open_document_encryption_tool,style="SmallRedButton.TButton").pack(pady=10)
        tb.Button(main_frame, text="Simulate Cloud Upload", command=simulate_upload, style="SmallRedButton.TButton").pack(pady=10)
        

    else:
        # Buttons for decryption mode
        tb.Button(main_frame, text="Simulate Cloud Download", command=simulate_upload, style="SmallRedButton.TButton").pack(pady=10)
        tb.Button(main_frame, text="Request Data Access", command=lambda: request_access_window(current_user_id), style="SmallRedButton.TButton").pack(pady=10)
        tb.Button(main_frame, text="Upload File", command=upload_file, style="SmallRedButton.TButton").pack(pady=10)
        tb.Button(main_frame, text="Decrypt Data", command=decrypt_data, style="SmallRedButton.TButton").pack(pady=10)

   
    # 🔙 Return Button
    tb.Button(main_frame, text="Return", command=lambda: return_to_operation_choice(main_frame), bootstyle="danger-outline").pack(pady=10)

    # Add Sign Out button
    tb.Button(main_frame, text="Sign Out", command=create_signin_page, bootstyle="danger-outline").pack(side="top", anchor="nw", padx=10, pady=10)

def create_signup_page():
    for widget in root.winfo_children():
        widget.destroy()

    # Create a canvas and a scrollbar for scrolling functionality
    canvas = tb.Canvas(root)  # Removed bootstyle from here
    scrollbar = tb.Scrollbar(root, orient="vertical", command=canvas.yview)
    scrollable_frame = tb.Frame(canvas, padding=(20, 10), bootstyle="white")

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Organize widgets within the scrollable_frame
    # Add logo
    add_logo(scrollable_frame)

    def create_signup_page():
        for widget in root.winfo_children():
            widget.destroy()

        # Create a canvas and a scrollbar for scrolling functionality
        canvas = tb.Canvas(root)  # Removed bootstyle from here
        scrollbar = tb.Scrollbar(root, orient="vertical", command=canvas.yview)
        scrollable_frame = tb.Frame(canvas, padding=(20, 10), bootstyle="white")

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Organize widgets within the scrollable_frame
        # Add logo
        add_logo(scrollable_frame)

    def submit_signup():
        first_name = entry_first_name.get()
        last_name = entry_last_name.get()
        email = entry_email.get()
        password = entry_password.get()
        confirm_password = entry_confirm_password.get()
        department = department_var.get()
        security_question = security_question_var.get()
        security_answer = entry_security_answer.get()

        if password != confirm_password:
            messagebox.showerror("Error", "Passwords do not match.")
            return

        hashed_password = hash_password(password)

        try:
            conn = sqlite3.connect('users.db')
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO users (email, password, first_name, last_name, department, security_question, security_answer)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (email, hashed_password, first_name, last_name, department, security_question, security_answer))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "User registered successfully!")
        except sqlite3.IntegrityError:
            messagebox.showerror("Error", "User already exists.")

        create_signin_page()

    # Organizing fields using `pack()` to avoid mix of geometry managers
    tb.Label(scrollable_frame, text="First Name:", bootstyle="info").pack(anchor="w", padx=10, pady=5)
    entry_first_name = tb.Entry(scrollable_frame)
    entry_first_name.pack(fill="x", padx=10, pady=5)

    tb.Label(scrollable_frame, text="Last Name:", bootstyle="info").pack(anchor="w", padx=10, pady=5)
    entry_last_name = tb.Entry(scrollable_frame)
    entry_last_name.pack(fill="x", padx=10, pady=5)

    tb.Label(scrollable_frame, text="Email:", bootstyle="info").pack(anchor="w", padx=10, pady=5)
    entry_email = tb.Entry(scrollable_frame)
    entry_email.pack(fill="x", padx=10, pady=5)

    tb.Label(scrollable_frame, text="Password:", bootstyle="info").pack(anchor="w", padx=10, pady=5)
    entry_password = tb.Entry(scrollable_frame, show="*")
    entry_password.pack(fill="x", padx=10, pady=5)

    tb.Label(scrollable_frame, text="Confirm Password:", bootstyle="info").pack(anchor="w", padx=10, pady=5)
    entry_confirm_password = tb.Entry(scrollable_frame, show="*")
    entry_confirm_password.pack(fill="x", padx=10, pady=5)

    tb.Label(scrollable_frame, text="Department:", bootstyle="info").pack(anchor="w", padx=10, pady=5)
    department_var = tb.StringVar(value="Marketing")
    departments = [("Marketing", "Marketing"), ("IT", "IT"), ("Infrastructure", "Infrastructure"), ("Admin", "Admin")]
    for text, value in departments:
        tb.Radiobutton(scrollable_frame, text=text, variable=department_var, value=value, bootstyle="danger-toolbutton").pack(anchor="w", padx=10, pady=5)

    # Security Question and Answer
    security_question_var = tb.StringVar(value="Which is your favorite movie?")
    security_questions = [
        "Which is your favorite movie?",
        "What was your first pet's name?",
        "What is your mother's maiden name?",
        "What is the name of your first school?",
        "What is your favorite food?"
    ]

    tb.Label(scrollable_frame, text="Security Question:", bootstyle="info").pack(anchor="w", padx=10, pady=5)
    question_menu = tb.OptionMenu(scrollable_frame, security_question_var, *security_questions)
    question_menu.pack(fill="x", padx=10, pady=5)

    tb.Label(scrollable_frame, text="Answer:", bootstyle="info").pack(anchor="w", padx=10, pady=5)
    entry_security_answer = tb.Entry(scrollable_frame)
    entry_security_answer.pack(fill="x", padx=10, pady=5)

    tb.Button(scrollable_frame, text="Sign Up", command=submit_signup, bootstyle="success").pack(fill="x", pady=20)

    # Return button to sign-in page
    tb.Button(scrollable_frame, text="Return", command=create_signin_page, bootstyle="secondary-outline").pack(fill="x", pady=10)

def create_signin_page():
    global entry_email, entry_password  # Declare widgets globally if needed

    # Clear any existing widgets from the root window
    for widget in root.winfo_children():
        widget.destroy()

    # Create a new frame for the sign-in page
    signin_frame = tb.Frame(root, padding=(20, 10), bootstyle="white")
    signin_frame.pack(fill="both", expand=True)

    add_logo(signin_frame)  # Assuming there's a function to add the logo

    # Create labels and entry widgets for the sign-in page
    tb.Label(signin_frame, text="Email:", bootstyle="info").pack(pady=5)
    entry_email = tb.Entry(signin_frame)
    entry_email.pack(pady=5)

    tb.Label(signin_frame, text="Password:", bootstyle="info").pack(pady=5)
    entry_password = tb.Entry(signin_frame, show="*")
    entry_password.pack(pady=5)

    # Sign In button that triggers the sign-in logic
    tb.Button(signin_frame, text="Sign In", command=submit_signin, style="SmallRedButton.TButton").pack(pady=20)

    # Forgot password button
    tb.Button(signin_frame, text="Forgot Password?", command=forgot_password, style="WhiteButton.TButton").pack(pady=5)

    # Sign-up button
    tb.Button(signin_frame, text="Sign Up", command=create_signup_page, style="SmallRedButton.TButton").pack(pady=5)

# Global variables at module-level
current_user_id = None
current_admin_id = None
current_user_department = None
current_user_email = None


def submit_signin():
    global current_user_id, current_admin_id, current_user_department, current_user_email

    email = entry_email.get().strip()
    password = entry_password.get().strip()
    hashed_password = hash_password(password)

    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()
    cursor.execute('SELECT user_id, password, department FROM users WHERE email = ?', (email,))
    result = cursor.fetchone()
    conn.close()

    if result:
        db_user_id, db_password, db_department = result
        if db_password == hashed_password:
            if db_department.lower() == 'admin':
                current_admin_id = db_user_id
                current_user_id = db_user_id
                current_user_department = db_department
                current_user_email = email
                messagebox.showinfo("Success", "Admin login successful!")
            else:
                current_user_id = db_user_id
                current_user_department = db_department
                current_user_email = email
                messagebox.showinfo("Success", "Sign-in successful!")
            
            create_operation_choice_page()  # proceed to next screen
        else:
            messagebox.showerror("Error", "Invalid email or password.")
    else:
        messagebox.showerror("Error", "Invalid email or password.")




# Function to navigate to the operation choice page
def create_operation_choice_page():
    for widget in root.winfo_children():
        widget.destroy()

    choice_frame = tb.Frame(root, padding=(20, 10), bootstyle="white")
    choice_frame.pack(fill="both", expand=True)

    tb.Label(choice_frame, text="Choose an Operation", font=("Arial", 14, "bold"), bootstyle="primary").pack(pady=20)

    # Encryption & Decryption
    tb.Button(choice_frame, text="Encryption", command=lambda: create_main_page(is_encryption=True), bootstyle="info-outline").pack(fill="x", padx=50, pady=10)
    tb.Button(choice_frame, text="Decryption", command=lambda: create_main_page(is_encryption=False), bootstyle="info-outline").pack(fill="x", padx=50, pady=10)

    # Equipment Management
    tb.Button(choice_frame, text="Equipment Management", command=manage_equipment, bootstyle="success-outline").pack(fill="x", padx=50, pady=10)

    # Project Management
    tb.Button(choice_frame, text="Project Management", command=manage_projects, bootstyle="warning-outline").pack(fill="x", padx=50, pady=10)

    # Admin Panel
    tb.Button(choice_frame, text="Admin panel", command=show_dashboard, bootstyle="danger-outline").pack(fill="x", padx=50, pady=10)

    # Sign Out
    tb.Button(choice_frame, text="Sign Out", command=create_signin_page, bootstyle="secondary-outline").pack(fill="x", padx=50, pady=20)


# Call this function when the application starts to show the sign-in page
create_signin_page()
()




root.mainloop()

